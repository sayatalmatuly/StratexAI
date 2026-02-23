import io
import json
import os
import re
from datetime import datetime
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse

import PyPDF2
import requests
from docx import Document
from flask import Flask, jsonify, request, send_from_directory
from flask_cors import CORS
from openai import OpenAI
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy import Column, DateTime, Integer, String, Text, create_engine, text
from sqlalchemy.exc import OperationalError
from sqlalchemy.orm import declarative_base, sessionmaker

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

GENERATED_DIR = os.path.join(BASE_DIR, "generated")

os.makedirs(GENERATED_DIR, exist_ok=True)

def markdown_to_html(text: str) -> str:

    if not text:

        return ""

    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)

    text = re.sub(r"\*(.*?)\*", r"<em>\1</em>", text)

    text = re.sub(r"`([^`]*)`", r"\1", text)

    text = re.sub(r"^#{1,6}\s*(.*)$", r"<h3>\1</h3>", text, flags=re.MULTILINE)

    text = re.sub(r"\n{2,}", "<br><br>", text)

    return text

def generate_excel(text: str, filename: str):

    wb = Workbook()

    ws = wb.active

    ws.title = "Report"

    def parse_line(line: str):

        line = (line or "").strip()

        if not line:

            return None

        if re.match(r"^[\s\-\|]+$", line):

            return "sep", []

        if "\t" in line:

            return "row", [p.strip() for p in line.split("\t")]

        if "|" in line:

            parts = [p.strip() for p in re.split(r"\s*\|\s*", line)]

            parts = [p for p in parts if p != ""]

            if len(parts) == 1:

                return "header", parts

            if parts:

                return "row", parts

        if "; " in line:

            parts = [p.strip() for p in line.split("; ")]

            if len(parts) >= 2:

                return "row", parts

        if ": " in line:

            left, right = line.split(": ", 1)

            if len(left) <= 40:

                return "row", [left.strip(), right.strip()]

        line = re.sub(r"^[-•]\s*", "", line)

        return "header", [line]

    blocks = []

    current = []

    for raw in text.split("\n"):

        parsed = parse_line(raw)

        if parsed is None:

            if current:

                blocks.append(current)

                current = []

            continue

        kind, parts = parsed

        if kind == "sep":

            continue

        current.append((kind, parts))

    if current:

        blocks.append(current)

    row_idx = 1

    header_font = Font(bold=True)

    for block in blocks:

        max_cols = max((len(p) for k, p in block if k == "row"), default=1)

        for kind, parts in block:

            if kind == "header":

                ws.cell(row=row_idx, column=1, value=parts[0]).font = header_font

                if max_cols > 1:

                    ws.merge_cells(
                        start_row=row_idx,
                        start_column=1,
                        end_row=row_idx,
                        end_column=max_cols,
                    )

                row_idx += 1

                continue

            for j, part in enumerate(parts, start=1):

                ws.cell(row=row_idx, column=j, value=part)

            row_idx += 1

        row_idx += 1

    path = os.path.join(GENERATED_DIR, f"{filename}.xlsx")

    wb.save(path)

    return path

def generate_word(text: str, filename: str):

    doc = Document()

    for line in text.split("\n"):

        doc.add_paragraph(line)

    path = os.path.join(GENERATED_DIR, f"{filename}.docx")

    doc.save(path)

    return path

def generate_pdf(text: str, filename: str):

    path = os.path.join(GENERATED_DIR, f"{filename}.pdf")

    c = canvas.Canvas(path, pagesize=A4)

    y = 800

    for line in text.split("\n"):

        c.drawString(40, y, line[:100])

        y -= 14

        if y < 40:

            c.showPage()

            y = 800

    c.save()

    return path

def strip_html(text: str) -> str:

    if not text:

        return ""

    text = re.sub(r"<br\s*\/?>", "\n", text)

    text = re.sub(r"<[^>]*>", "", text)

    return text.strip()

def clean_document_content(text: str) -> str:

    if not text:

        return ""

    lines = [l.strip() for l in text.splitlines()]

    cleaned = []

    for line in lines:

        if not line:

            cleaned.append(line)

            continue

        l = line.lower()

        if (
            l.startswith("download ")
            or "download the pdf" in l
            or "download the docx" in l
            or "download the word" in l
        ):

            continue

        if (
            l.startswith("скачать ")
            or "скачать pdf" in l
            or "скачать docx" in l
            or "скачать word" in l
            or "скачать excel" in l
        ):

            continue

        if l.startswith("here is") and (
            "format" in l or "document" in l or "pdf" in l or "docx" in l or "word" in l
        ):

            continue

        if l.startswith("here is the") and (
            "format" in l or "document" in l or "pdf" in l or "docx" in l or "word" in l
        ):

            continue

        if l.startswith("вот") and (
            "документ" in l or "pdf" in l or "docx" in l or "word" in l or "excel" in l
        ):

            continue

        if l.startswith("ваш документ") or "документ готов" in l:

            continue

        if "ссылка" in l and ("документ" in l or "скачать" in l):

            continue

        cleaned.append(line)

    out = "\n".join(cleaned)

    out = re.sub(r"\n{3,}", "\n\n", out).strip()

    return out

def dedupe_document_content(text: str) -> str:

    if not text:

        return ""

    lines = [l.rstrip() for l in text.splitlines()]

    paragraphs = []

    buf = []

    for line in lines:

        if line.strip() == "":

            if buf:

                paragraphs.append("\n".join(buf))

                buf = []

        else:

            buf.append(line)

    if buf:

        paragraphs.append("\n".join(buf))

    deduped = []

    prev_norm = None

    for p in paragraphs:

        norm = " ".join(p.split()).strip().lower()

        if prev_norm is not None and norm == prev_norm:

            continue

        deduped.append(p)

        prev_norm = norm

    return "\n\n".join(deduped).strip()


def _ensure_sslmode_require_for_prod(database_url: str) -> str:
    """
    Ensure sslmode=require for non-local PostgreSQL URLs (for production hosting).
    Does not force SSL for localhost/127.0.0.1 to keep local development simple.
    """
    try:
        parsed = urlparse(database_url)
        if not parsed.scheme.startswith("postgresql"):
            return database_url

        host = (parsed.hostname or "").lower()
        if host in ("localhost", "127.0.0.1"):
            return database_url

        query_params = dict(parse_qsl(parsed.query, keep_blank_values=True))
        if "sslmode" not in query_params:
            query_params["sslmode"] = "require"
            parsed = parsed._replace(query=urlencode(query_params))
            return urlunparse(parsed)

        return database_url
    except Exception:
        # In case of any parsing issues, fall back to the original URL
        return database_url

def detect_file_type(message: str) -> str:

    m = (message or "").lower()

    if "xlsx" in m or "xls" in m or "excel" in m or "эксель" in m or "ексель" in m:

        return "excel"

    if "docx" in m or "doc" in m or "word" in m or "ворд" in m or "док" in m:

        return "word"

    if "pdf" in m or "пдф" in m:

        return "pdf"

    return "pdf"

def explicit_file_request(message: str) -> bool:
    """
    Check if user explicitly requests a file download/generation.
    Only returns True if user uses action verbs like "create", "generate", "download", "save".
    """
    m = (message or "").lower()

    # Action verbs that indicate file generation request
    action_keywords = [
        "create", "generate", "download", "save", "export", "make", "build",
        "создай", "создать", "сгенерируй", "скачай", "скачать", "сохрани", "сохранить",
        "экспорт", "экспортировать", "сделай", "сделать", "построй", "построить"
    ]

    # File format keywords
    file_keywords = [
        "pdf", "docx", "doc", "word", "xlsx", "xls", "excel",
        "пдф", "ворд", "эксель", "документ", "файл"
    ]

    # Check if message contains both action verb AND file keyword
    has_action = any(action in m for action in action_keywords)
    has_file = any(keyword in m for keyword in file_keywords)

    # Only return True if BOTH action verb and file keyword are present
    return has_action and has_file


def _test_connection(db_engine):
    with db_engine.connect() as conn:
        conn.execute(text("SELECT 1"))


DATABASE_URL_PRIVATE_RAW = os.getenv(
    "DATABASE_URL", "postgresql+psycopg2://postgres:alm@localhost:5432/stratex"
)
DATABASE_URL_PUBLIC_RAW = os.getenv("DATABASE_PUBLIC_URL")

DATABASE_URL_PRIVATE = (
    _ensure_sslmode_require_for_prod(DATABASE_URL_PRIVATE_RAW)
    if DATABASE_URL_PRIVATE_RAW
    else None
)
DATABASE_URL_PUBLIC = (
    _ensure_sslmode_require_for_prod(DATABASE_URL_PUBLIC_RAW)
    if DATABASE_URL_PUBLIC_RAW
    else None
)

DATABASE_URL = DATABASE_URL_PRIVATE or DATABASE_URL_PUBLIC
if not DATABASE_URL:
    raise RuntimeError("No PostgreSQL URL configured. Set DATABASE_URL or DATABASE_PUBLIC_URL.")

engine = create_engine(
    DATABASE_URL,
    echo=False,
    future=True,
    pool_pre_ping=True,
    pool_recycle=300,
)

SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)

Base = declarative_base()

_CURRENT_DB_TARGET = "private" if DATABASE_URL == DATABASE_URL_PRIVATE else "public"

class ChatMessage(Base):

    __tablename__ = "chat_messages"

    id = Column(Integer, primary_key=True)

    user_id = Column(String, index=True)

    role = Column(String)

    content = Column(Text)

    created_at = Column(DateTime, default=datetime.utcnow)

Base.metadata.create_all(engine)

try:
    _test_connection(engine)
except OperationalError as e:
    print(f"[db] Initial connection to { _CURRENT_DB_TARGET } database failed: {e}")
    if DATABASE_URL_PUBLIC and _CURRENT_DB_TARGET == "private":
        print("[db] Trying to fall back to public database URL")
        fallback_engine = create_engine(
            DATABASE_URL_PUBLIC,
            echo=False,
            future=True,
            pool_pre_ping=True,
            pool_recycle=300,
        )
        _test_connection(fallback_engine)
        engine.dispose()
        engine = fallback_engine
        SessionLocal.configure(bind=engine)
        DATABASE_URL = DATABASE_URL_PUBLIC
        _CURRENT_DB_TARGET = "public"
        Base.metadata.create_all(engine)
    else:
        raise


def _switch_engine_to_public_on_error(error: OperationalError) -> bool:
    """
    On OperationalError against the private DB, switch engine to the public URL if available.
    Returns True when the switch succeeds, False otherwise.
    """
    global engine, DATABASE_URL, _CURRENT_DB_TARGET

    if not DATABASE_URL_PUBLIC or _CURRENT_DB_TARGET == "public":
        return False

    print(f"[db] OperationalError on private DB, switching to public: {error}")
    try:
        new_engine = create_engine(
            DATABASE_URL_PUBLIC,
            echo=False,
            future=True,
            pool_pre_ping=True,
            pool_recycle=300,
        )
        _test_connection(new_engine)
        engine.dispose()
        engine = new_engine
        SessionLocal.configure(bind=engine)
        DATABASE_URL = DATABASE_URL_PUBLIC
        _CURRENT_DB_TARGET = "public"
        Base.metadata.create_all(engine)
        return True
    except OperationalError as e:
        print(f"[db] Fallback to public DB failed: {e}")
        return False


def _run_db_operation(operation):
    """
    Run a DB operation with automatic fallback from private to public database on OperationalError.
    """
    try:
        return operation()
    except OperationalError as e:
        if _switch_engine_to_public_on_error(e):
            return operation()
        raise

def db_add_message(user_id: str, role: str, content: str):

    def _op():
        db = SessionLocal()
        try:
            db.add(ChatMessage(user_id=user_id, role=role, content=content))
            db.commit()
        finally:
            db.close()

    _run_db_operation(_op)

def db_get_history(user_id: str, limit: int = 20):

    def _op():
        db = SessionLocal()
        try:
            rows = (
                db.query(ChatMessage)
                .filter(ChatMessage.user_id == user_id)
                .order_by(ChatMessage.created_at.asc())
                .all()
            )
            rows_limited = rows[-limit:]
            return [{"role": r.role, "content": r.content} for r in rows_limited]
        finally:
            db.close()

    return _run_db_operation(_op)

app = Flask(__name__)

# Configure CORS to allow all origins (including file:// and null origin)
CORS(app, resources={
    r"/api/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

@app.route("/api/generate-file", methods=["POST"])
def generate_file():

    data = request.get_json(silent=True) or {}

    content = data.get("content", "")

    file_type = data.get("file_type")

    if (not content) or (file_type not in ["excel", "word", "pdf"]):

        return jsonify({"success": False, "error": "Invalid request"}), 400

    name = f"report_{int(datetime.now().timestamp())}"

    if file_type == "excel":

        path = generate_excel(content, name)

    elif file_type == "word":

        path = generate_word(content, name)

    else:

        path = generate_pdf(content, name)

    try:

        size = os.path.getsize(path)

    except Exception:

        size = -1

    print(f"[generate_file] path={path} exists={os.path.exists(path)} size={size}")

    return jsonify(
        {"success": True, "download_url": f"/api/download/{os.path.basename(path)}"}
    )

@app.route("/api/download/<filename>", methods=["GET"])
def download_file(filename):

    path = os.path.join(GENERATED_DIR, filename)

    try:

        size = os.path.getsize(path)

    except Exception:

        size = -1

    print(f"[download_file] path={path} exists={os.path.exists(path)} size={size}")

    return send_from_directory(GENERATED_DIR, filename, as_attachment=True)

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")

openai_client = OpenAI(api_key=OPENAI_API_KEY)

USERS = {"demo": "demo123", "business": "business2024"}

# Web search via Tavily

def web_search(query):

    if not TAVILY_API_KEY:

        return {"answer": "", "sources": [], "error": "TAVILY_API_KEY is not set"}

    try:

        url = "https://api.tavily.com/search"

        payload = {
            "api_key": TAVILY_API_KEY,
            "query": query,
            "search_depth": "advanced",
            "include_answer": True,
            "max_results": 5,
        }

        response = requests.post(url, json=payload)

        data = response.json()

        results = {"answer": data.get("answer", ""), "sources": []}

        for result in data.get("results", []):

            results["sources"].append(
                {
                    "title": result.get("title", ""),
                    "url": result.get("url", ""),
                    "content": result.get("content", "")[:300],
                }
            )

        return results

    except Exception as e:

        print(f"Search error: {e}")

        return {"answer": "", "sources": []}

def extract_pdf_text(file_bytes):

    try:

        pdf_file = io.BytesIO(file_bytes)

        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text = ""

        for page in pdf_reader.pages:

            text += page.extract_text()

        return text

    except Exception as e:

        return f"PDF reading error: {str(e)}"

def extract_docx_text(file_bytes):

    try:

        doc_file = io.BytesIO(file_bytes)

        doc = Document(doc_file)

        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        return text

    except Exception as e:

        return f"DOCX reading error: {str(e)}"

def detect_language(message: str) -> str:

    if re.search(r"[\u0400-\u04FF]", message or ""):

        return "ru"

    return "en"

def chat_with_ai(
    message, system_prompt, conversation_history=None, search_context=None
):

    try:

        if not OPENAI_API_KEY:

            return "OpenAI API key is not configured. Please set OPENAI_API_KEY."

        messages = []

        lang = detect_language(message)

        lang_name = "Russian" if lang == "ru" else "English"

        full_system_prompt = (
            system_prompt + "\n\n"
            f"Current date: {datetime.now().strftime('%d.%m.%Y')}\n\n"
            f"Language: {lang_name}. Reply strictly in {lang_name}.\n\n"
            "Hard rule: If web search is required, silently perform it and then answer. "
            "Never say you will search, are searching, or will answer later. "
            "No meta-commentary about searching.\n\n"
            "Formatting rules:\n"
            "- DO NOT use Markdown (** ## `).\n"
            "- Use ONLY simple HTML tags.\n"
            "- Allowed tags: <strong>, <em>, <u>, <br>, <p>, <ul>, <li>, <h3>\n"
            "- Do NOT wrap everything in tags.\n"
        )

        if search_context:

            full_system_prompt += (
                "\n\nContext from the web:\n"
                f"{search_context}\n\n"
                "Hard rule: Use ONLY the web context above. Do NOT add facts not present there. "
                "If the context is insufficient, say so and ask to уточнить запрос."
            )

        messages.append({"role": "system", "content": full_system_prompt})

        if conversation_history:

            messages.extend(conversation_history[-10:])

        messages.append({"role": "user", "content": message})

        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=messages,
            temperature=0.7,
            max_tokens=2048,
            top_p=0.9,
        )

        return response.choices[0].message.content

    except Exception as e:

        print(f"AI Error: {e}")

        return f"An error occurred while processing the request: {str(e)}"

def detect_intent(message):

    message_lower = (message or "").lower()

    search_keywords = [
        "market",
        "trend",
        "news",
        "relevant",
        "price",
        "rate",
        "latest",
        "today",
        "now",
        "current",
        "search",
        "browse",
        "web",
        "internet",
        "online",
        "source",
        "sources",
        "salary",
        "wage",
        "pay",
        "compensation",
        "find",
        "look up",
        "check",
        "verify",
        "research",
        "?????",
        "?????",
        "??????",
        "????",
        "????",
        "???????",
        "??????",
        "????????",
        "?????",
        "?????",
        "???????",
        "?????",
        "???????",
        "???????",
        "?????????",
        "?????",
        "??????",
        "????????",
        "? ????",
        "? ?????????",
        "????????",
        "??????",
        "??? ?????",
        "??? ?????",
        "??? ?????",
        "??? ??",
        "????? ?????????",
        "??????? ????????",
        "former name",
        "???????",
        "?????",
        "?????",
    ]

    needs_search = any(keyword in message_lower for keyword in search_keywords)
    if not needs_search:
        if re.search(
            r"\b(кто|что|когда|где|почему|зачем|как|чей|чья|чьё|чьи)\b", message_lower
        ):
            needs_search = True
        elif re.search(r"\b(who|what|when|where|why|how|whose)\b", message_lower):
            needs_search = True
        elif "?" in message_lower:
            needs_search = True
    doc_keywords = [
        "report",
        "plan",
        "presentation",
        "proposal",
        "document",
        "email",
        "e-mail",
        "letter",
        "?????",
        "????",
        "?????????",
        "?????????",
        "????????",
        "??????",
        "?????",
        "e-mail",
    ]

    file_keywords = [
        "docx",
        "doc",
        "word",
        "pdf",
        "xlsx",
        "xls",
        "excel",
        "????",
        "???",
        "???",
        "?????",
        "excel",
    ]

    needs_document = any(keyword in message_lower for keyword in doc_keywords) or any(
        keyword in message_lower for keyword in file_keywords
    )

    analyze_keywords = [
        "analyze",
        "check",
        "evaluate",
        "review",
        "study",
        "??????",
        "???????",
        "????",
        "??????",
    ]

    needs_analysis = any(keyword in message_lower for keyword in analyze_keywords)

    return {
        "needs_search": needs_search,
        "needs_document": needs_document,
        "needs_analysis": needs_analysis,
    }

@app.route("/api/auth", methods=["POST"])
def authenticate():

    data = request.json

    username = data.get("username", "")

    password = data.get("password", "")

    if username in USERS and USERS[username] == password:

        return jsonify(
            {
                "success": True,
                "user": {
                    "username": username,
                    "tier": "premium" if username == "business" else "free",
                },
            }
        )

    return jsonify({"success": False, "error": "Invalid username or password"}), 401

@app.route("/api/chat", methods=["POST"])
def chat():

    try:

        data = request.json

        message = data.get("message", "")

        user_id = data.get("user_id", "guest")

        is_guest = data.get("is_guest", True)

        # Флаг принудительного веб-поиска с фронтенда
        force_web_search = bool(data.get("force_web_search"))

        intent = detect_intent(message)

        # Если тумблер Web Search включён на фронтенде, всегда считаем,
        # что нужен поиск (остальная логика использует intent["needs_search"])
        if force_web_search:
            intent["needs_search"] = True

        if is_guest:

            if (
                intent["needs_search"]
                or intent["needs_document"]
                or intent["needs_analysis"]
            ):

                return jsonify(
                    {
                        "success": True,
                        "message": "??? ??????? ???????? ?????? ????? ?????. ??????? ?Login? ? ??????? ? ???????.",
                        "require_auth": True,
                    }
                )

            if len(message) > 500:

                return jsonify(
                    {
                        "success": False,
                        "message": "Guest mode is limited to 500 characters. Please log in for full access.",
                        "require_auth": True,
                    }
                )

            system_prompt = """You are the demo version of Stratex AI. Reply briefly and suggest logging in to unlock:

- Market and trend analysis

- Full report generation

- Working with documents

- Advanced features

Use context from previous messages when relevant."""

            conversation_history = db_get_history(user_id, limit=10)

            response = chat_with_ai(message, system_prompt, conversation_history)

            db_add_message(user_id, "user", message)

            db_add_message(user_id, "assistant", response)

            return jsonify(
                {
                    "success": True,
                    "message": response,
                    "is_demo": True,
                    "upgrade_prompt": "?? Log in for full access",
                }
            )

        conversation_history = db_get_history(user_id, limit=20)

        search_context = None

        search_results = None

        if intent["needs_search"]:

            search_data = web_search(message)

            if search_data.get("error"):
                # Do not fail the whole chat when web search is unavailable.
                # This can happen when TAVILY_API_KEY is not set or web search fails.
                # We simply continue without search context.
                print("[chat] Web search unavailable or error; continuing without web context")
                intent["needs_search"] = False
                search_context = None
                search_results = None
            else:
                sources = search_data.get("sources", [])

                if not sources:
                    # Do not fail the whole chat when web search yields no sources.
                    # This case can happen during document workflows where the message
                    # accidentally triggers needs_search, or when Tavily returns no results.
                    # We simply continue without search context.
                    print("[chat] Web search returned 0 sources; continuing without web context")
                    intent["needs_search"] = False
                    search_context = None
                    search_results = None
                else:
                    # Only build search context if we have sources
                    answer = search_data.get("answer", "")

                    search_context = f"Up-to-date information:\n{answer}\n\nSources:\n"

                    for source in sources:

                        search_context += (
                            f"- {source.get('title','')}: {source.get('content','')}\n"
                        )

                    search_results = search_data

        system_prompt = """You are Stratex AI, a professional business assistant.

Your specializations:

?? Market and trend analysis (with web search)

?? Report, plan, and document generation

?? Business communication (emails, proposals)

?? Strategic planning

Working principles:

1. Structured output ? use headings, lists, and tables

2. Measurable ? provide concrete KPIs, timelines, numbers

3. Professional ? business tone

4. Practical ? focus on execution

5. Use context from previous messages when relevant

Table formatting:

For tabular data, use a simple pipe-separated table (plain text, no markdown styling):

Header 1 | Header 2

--- | ---

Row 1 | Row 2

Document formats:

If the user asks for a report/plan ? FIRST ask for the format:

- Brief (1 page)

- Standard (2?3 pages)

- Detailed (5+ pages)

- Presentation (slides)

After answering ? generate in the chosen format."""

        if intent["needs_search"] and not search_context:

            system_prompt += "\n\nImportant: If the user asks for up-to-date web info and no web context is provided, explicitly say that live web search is unavailable right now and ask to enable TAVILY_API_KEY."

        ai_response = chat_with_ai(
            message, system_prompt, conversation_history, search_context
        )

        ai_response = markdown_to_html(ai_response)

        # Don't generate file if this is a document analysis response
        # Check if message contains document analysis pattern (starts with "Document" or contains analysis markers)
        is_document_analysis_response = (
            message.startswith("Document") or 
            message.startswith("Документ") or
            "Please use these document analyses" in message or
            "Document \"" in message
        )

        if intent.get("needs_document") and explicit_file_request(message) and not is_document_analysis_response:

            try:

                file_type = detect_file_type(message)

                name = f"report_{int(datetime.now().timestamp())}"

                doc_prompt = (
                    "Write the full document content as plain text based on the user's request. "
                    "Do not mention files, formats, or that you cannot create documents. "
                    "Do not add prefaces, links, or download instructions. "
                    "Output only the document content."
                )

                doc_text = chat_with_ai(message, doc_prompt, conversation_history)

                content = dedupe_document_content(
                    clean_document_content(strip_html(doc_text))
                ).strip()

                if len(content) < 50:

                    raise ValueError("Generated content too short for document export")

                if file_type == "excel":

                    path = generate_excel(content, name)

                elif file_type == "word":

                    path = generate_word(content, name)

                else:

                    path = generate_pdf(content, name)

                base_url = request.host_url.rstrip("/")

                # Only convert http to https for production (not localhost)
                # For localhost, keep http:// as it won't work with https
                if base_url.startswith("http://") and "localhost" not in base_url and "127.0.0.1" not in base_url:
                    base_url = "https://" + base_url[len("http://") :]

                download_url = f"{base_url}/api/download/{os.path.basename(path)}"

                doc_label = {"word": "DOCX", "excel": "XLSX", "pdf": "PDF"}.get(
                    file_type, "DOC"
                )

                link_label = {"word": "WORD", "excel": "EXCEL", "pdf": "PDF"}.get(
                    file_type, "DOC"
                )

                lang = detect_language(message)

                if lang == "ru":

                    title = f"Ваш документ {doc_label} готов:"

                    link_text = f"Скачать {link_label}"

                else:

                    title = f"Here is your {doc_label} document:"

                    link_text = f"Download {link_label}"

                ai_response = (
                    f"{title}<br>"
                    f'<a href="{download_url}" target="_blank" '
                    f'rel="noopener" style="color:#6f47eb;font-weight:600;">'
                    f"{link_text}</a>"
                )

            except Exception as e:

                print(f"File generation error: {e}")

        db_add_message(user_id, "user", message)

        db_add_message(user_id, "assistant", ai_response)

        response_data = {
            "success": True,
            "message": ai_response,
            "is_demo": False,
            "debug": {
                "needs_search": bool(intent.get("needs_search")),
                "sources_count": (
                    len(search_results.get("sources", [])) if search_results else 0
                ),
            },
        }
        if search_results:

            response_data["sources"] = search_results["sources"]

        return jsonify(response_data)

    except Exception as e:

        print(f"Chat error: {e}")

        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/analyze-document", methods=["POST"])
def analyze_document():

    try:

        user_id = request.form.get("user_id", "guest")

        is_guest = request.form.get("is_guest", "true") == "true"

        if is_guest:

            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Document analysis is available only for authenticated users",
                        "require_auth": True,
                    }
                ),
                403,
            )

        if "file" not in request.files:

            return jsonify({"success": False, "error": "No file uploaded"}), 400

        file = request.files["file"]

        filename = file.filename

        file_bytes = file.read()

        if filename.endswith(".pdf"):

            text = extract_pdf_text(file_bytes)

        elif filename.endswith(".docx"):

            text = extract_docx_text(file_bytes)

        elif filename.endswith(".txt"):

            text = file_bytes.decode("utf-8")

        else:

            return jsonify({"success": False, "error": "Unsupported format"}), 400

        system_prompt = """You are an expert in business document analysis.

Analyze the document using these criteria:

1. Document type ? what is it?

2. Main content ? short summary

3. Key points ? 3?5 main points

4. Strengths ? what is done well

5. Areas for improvement ? what can be improved

6. Recommendations ? concrete next steps

Output format: structured with headings."""

        analysis = chat_with_ai(
            f"Analyze the document '{filename}':\n\n{text[:4000]}", system_prompt
        )

        return jsonify(
            {
                "success": True,
                "filename": filename,
                "analysis": analysis,
                "document_length": len(text),
            }
        )

    except Exception as e:

        print(f"Document analysis error: {e}")

        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/market-analysis", methods=["POST"])
def market_analysis():

    try:

        data = request.json

        query = data.get("query", "")

        is_guest = data.get("is_guest", True)

        if is_guest:

            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Market analysis is available only for authenticated users",
                        "require_auth": True,
                    }
                ),
                403,
            )

        search_results = web_search(query)

        system_prompt = """You are a market analyst. Based on the information found, create:

1. Current situation ? what is happening now

2. Key trends ? 3?5 key trends

3. Opportunities ? for businesses

4. Risks ? potential threats

5. Recommendations ? concrete actions

Use data from the sources and include numbers when possible."""

        context = f"Search results for '{query}':\n\n"

        for source in search_results.get("sources", []):

            context += f"- {source['title']}: {source['content']}\n"

        analysis = chat_with_ai(
            f"Analyze the market: {query}", system_prompt, search_context=context
        )

        return jsonify(
            {
                "success": True,
                "query": query,
                "analysis": analysis,
                "sources": search_results.get("sources", []),
            }
        )

    except Exception as e:

        print(f"Market analysis error: {e}")

        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/generate-report", methods=["POST"])
def generate_report():

    try:

        data = request.json

        report_type = data.get("type", "")

        parameters = data.get("parameters", {})

        format_type = data.get("format", "standard")

        is_guest = data.get("is_guest", True)

        if is_guest:

            return (
                jsonify(
                    {
                        "success": False,
                        "message": "Report generation is available only for authenticated users",
                        "require_auth": True,
                    }
                ),
                403,
            )

        format_instructions = {
            "brief": "1 page, key points as a list",
            "standard": "2?3 pages, structured with sections",
            "detailed": "5+ pages, detailed analysis with examples",
            "presentation": "presentation slide format",
        }

        system_prompt = f"""You are an expert in business documentation.

Create a {report_type} report in the following format: {format_instructions.get(format_type, 'standard')}

Required elements:

- Executive summary

- Structured headings

- Concrete numbers and KPIs

- Timelines

- Recommendations

Parameters: {json.dumps(parameters, ensure_ascii=False)}"""

        report = chat_with_ai(f"Create a {report_type} report", system_prompt)

        return jsonify(
            {
                "success": True,
                "report": report,
                "type": report_type,
                "format": format_type,
            }
        )

    except Exception as e:

        print(f"Report generation error: {e}")

        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/status", methods=["GET"])
def status():

    return jsonify(
        {
            "success": True,
            "status": "online",
            "openai": "configured" if OPENAI_API_KEY else "missing_key",
            "tavily": "configured" if TAVILY_API_KEY else "missing_key",
            "model": OPENAI_MODEL,
        }
    )

if __name__ == "__main__":

    print("Stratex AI Backend started!")

    print("OpenAI API: connected")

    print("Tavily Search: connected")

    print("Server: http://localhost:5000")

    print("Test accounts:")

    print("   demo / demo123 (basic)")

    print("   business / business2024 (full)")

    app.run(debug=True, port=5000, host="0.0.0.0")
